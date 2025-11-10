#!/usr/bin/env python3

"""
Export Markdown to Word Document with Rendered Mermaid Diagrams

This script converts the GitHub-ServiceNow integration documentation from
Markdown to Word format (.docx) with Mermaid diagrams rendered as images.

Requirements:
    pip install python-docx markdown beautifulsoup4 requests pillow

Optional (for better Mermaid rendering):
    npm install -g @mermaid-js/mermaid-cli
"""

import os
import sys
import re
import json
import subprocess
import tempfile
from pathlib import Path
from typing import List, Tuple
import base64
import io

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import markdown
    from bs4 import BeautifulSoup
except ImportError as e:
    print(f"ERROR: Missing required Python package: {e}")
    print("\nInstall required packages with:")
    print("  pip install python-docx markdown beautifulsoup4 requests pillow")
    sys.exit(1)


class MermaidRenderer:
    """Render Mermaid diagrams to PNG images"""

    def __init__(self):
        self.has_mmdc = self._check_mmdc()
        self.temp_dir = tempfile.mkdtemp()

    def _check_mmdc(self) -> bool:
        """Check if mermaid-cli (mmdc) is installed"""
        try:
            subprocess.run(['mmdc', '--version'],
                         capture_output=True, check=True)
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False

    def render_diagram(self, mermaid_code: str, diagram_num: int) -> str:
        """
        Render a Mermaid diagram to PNG

        Args:
            mermaid_code: Mermaid diagram source code
            diagram_num: Diagram number for file naming

        Returns:
            Path to rendered PNG image
        """
        if not self.has_mmdc:
            print(f"WARNING: mermaid-cli not installed, skipping diagram {diagram_num}")
            return None

        # Create temporary mermaid file
        mmd_file = os.path.join(self.temp_dir, f'diagram_{diagram_num}.mmd')
        png_file = os.path.join(self.temp_dir, f'diagram_{diagram_num}.png')

        # Write mermaid code to file
        with open(mmd_file, 'w', encoding='utf-8') as f:
            f.write(mermaid_code)

        # Render with mmdc
        try:
            subprocess.run(
                ['mmdc', '-i', mmd_file, '-o', png_file, '-b', 'transparent'],
                capture_output=True,
                check=True
            )
            print(f"✓ Rendered diagram {diagram_num}")
            return png_file
        except subprocess.CalledProcessError as e:
            print(f"WARNING: Failed to render diagram {diagram_num}: {e}")
            return None

    def cleanup(self):
        """Remove temporary files"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)


class MarkdownToWordConverter:
    """Convert Markdown with Mermaid diagrams to Word document"""

    def __init__(self, input_file: str, output_file: str):
        self.input_file = input_file
        self.output_file = output_file
        self.document = Document()
        self.renderer = MermaidRenderer()
        self.diagram_counter = 0

        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles"""
        # Set up custom styles for better formatting
        styles = self.document.styles

        # Title style
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(44, 62, 80)

        # Heading 1
        h1 = styles['Heading 1']
        h1.font.size = Pt(20)
        h1.font.color.rgb = RGBColor(44, 62, 80)

        # Heading 2
        h2 = styles['Heading 2']
        h2.font.size = Pt(16)
        h2.font.color.rgb = RGBColor(52, 73, 94)

        # Heading 3
        h3 = styles['Heading 3']
        h3.font.size = Pt(14)
        h3.font.color.rgb = RGBColor(52, 73, 94)

    def extract_mermaid_diagrams(self, content: str) -> Tuple[str, List[str]]:
        """
        Extract Mermaid diagram blocks from markdown

        Returns:
            Tuple of (modified_content, list_of_diagram_codes)
        """
        diagrams = []
        diagram_placeholder = "<<<DIAGRAM_{}>>>"

        def replace_diagram(match):
            diagram_code = match.group(1)
            diagrams.append(diagram_code)
            self.diagram_counter += 1
            return diagram_placeholder.format(self.diagram_counter)

        # Extract mermaid code blocks
        pattern = r'```mermaid\n(.*?)```'
        modified_content = re.sub(pattern, replace_diagram, content, flags=re.DOTALL)

        return modified_content, diagrams

    def add_title(self, text: str):
        """Add document title"""
        title = self.document.add_heading(text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_paragraph(self, text: str, style=None):
        """Add paragraph to document"""
        if not text.strip():
            return

        para = self.document.add_paragraph(text)
        if style:
            para.style = style

        return para

    def add_heading(self, text: str, level: int):
        """Add heading to document"""
        self.document.add_heading(text, level)

    def add_table(self, html_table):
        """Add table from HTML to document"""
        soup = BeautifulSoup(str(html_table), 'html.parser')

        # Get headers
        headers = [th.get_text(strip=True) for th in soup.find_all('th')]

        # Get rows
        rows = []
        for tr in soup.find_all('tr')[1:]:  # Skip header row
            cells = [td.get_text(strip=True) for td in tr.find_all('td')]
            if cells:
                rows.append(cells)

        if not headers or not rows:
            return

        # Create table
        table = self.document.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add rows
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell in enumerate(row):
                if j < len(row_cells):
                    row_cells[j].text = cell

        self.document.add_paragraph()  # Add spacing after table

    def add_code_block(self, code: str, language: str = None):
        """Add code block to document"""
        para = self.document.add_paragraph(code)
        para.style = 'Intense Quote'

        # Set monospace font
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    def add_diagram_image(self, diagram_code: str, diagram_num: int):
        """Add rendered Mermaid diagram as image"""
        image_path = self.renderer.render_diagram(diagram_code, diagram_num)

        if image_path and os.path.exists(image_path):
            # Add diagram caption
            caption = self.document.add_paragraph(f'Diagram {diagram_num}')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.bold = True
            caption.runs[0].font.size = Pt(10)

            # Add image
            try:
                self.document.add_picture(image_path, width=Inches(6.0))
                last_paragraph = self.document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"WARNING: Failed to add diagram {diagram_num} to document: {e}")

            # Add spacing
            self.document.add_paragraph()
        else:
            # Add placeholder text
            para = self.document.add_paragraph(
                f'[Diagram {diagram_num}: Mermaid diagram - install mermaid-cli to render]'
            )
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].italic = True

    def process_content(self, content: str):
        """Process markdown content and add to document"""
        # Extract diagrams first
        modified_content, diagrams = self.extract_mermaid_diagrams(content)

        # Convert markdown to HTML
        html = markdown.markdown(
            modified_content,
            extensions=['tables', 'fenced_code', 'codehilite']
        )

        soup = BeautifulSoup(html, 'html.parser')

        # Process each element
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'table', 'blockquote']):
            tag_name = element.name
            text = element.get_text(strip=True)

            # Check for diagram placeholder
            diagram_match = re.search(r'<<<DIAGRAM_(\d+)>>>', str(element))
            if diagram_match:
                diagram_num = int(diagram_match.group(1))
                if diagram_num <= len(diagrams):
                    self.add_diagram_image(diagrams[diagram_num - 1], diagram_num)
                continue

            if tag_name == 'h1':
                self.add_heading(text, 1)
            elif tag_name == 'h2':
                self.add_heading(text, 2)
            elif tag_name == 'h3':
                self.add_heading(text, 3)
            elif tag_name == 'h4':
                self.add_heading(text, 4)
            elif tag_name == 'p':
                self.add_paragraph(text)
            elif tag_name == 'pre':
                code = element.find('code')
                if code:
                    self.add_code_block(code.get_text())
                else:
                    self.add_code_block(text)
            elif tag_name == 'table':
                self.add_table(element)
            elif tag_name == 'blockquote':
                para = self.add_paragraph(text, 'Intense Quote')

    def convert(self):
        """Perform the conversion"""
        print(f"Reading input file: {self.input_file}")

        # Read markdown content
        with open(self.input_file, 'r', encoding='utf-8') as f:
            content = f.read()

        print("Processing content...")

        # Extract title
        title_match = re.match(r'^#\s+(.+)', content, re.MULTILINE)
        if title_match:
            self.add_title(title_match.group(1))
            # Remove title from content
            content = re.sub(r'^#\s+.+\n', '', content, count=1)

        # Process content
        self.process_content(content)

        # Save document
        print(f"Saving Word document: {self.output_file}")
        self.document.save(self.output_file)

        # Get file size
        file_size = os.path.getsize(self.output_file) / 1024  # KB
        print(f"✓ Document created successfully ({file_size:.1f} KB)")
        print(f"  Diagrams rendered: {self.diagram_counter}")

        # Cleanup
        self.renderer.cleanup()

        return self.output_file


def main():
    """Main entry point"""
    print("=" * 60)
    print("Markdown to Word Converter with Mermaid Diagram Support")
    print("=" * 60)
    print()

    # Determine file paths
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    docs_dir = project_root / 'docs'

    input_file = docs_dir / 'GITHUB-SERVICENOW-DATA-Integration-in-ARC.md'
    output_file = docs_dir / 'GITHUB-SERVICENOW-DATA-Integration-in-ARC.docx'

    # Check if input file exists
    if not input_file.exists():
        print(f"ERROR: Input file not found: {input_file}")
        sys.exit(1)

    # Convert
    converter = MarkdownToWordConverter(str(input_file), str(output_file))

    try:
        result = converter.convert()

        print()
        print("=" * 60)
        print("✓ Conversion completed successfully!")
        print("=" * 60)
        print()
        print(f"Output file: {result}")
        print()
        print("You can now open this file with:")
        print("  - Microsoft Word")
        print("  - LibreOffice Writer")
        print("  - Google Docs (upload)")

        if not converter.renderer.has_mmdc:
            print()
            print("NOTE: Mermaid diagrams were not rendered as images.")
            print("To enable diagram rendering, install mermaid-cli:")
            print("  npm install -g @mermaid-js/mermaid-cli")

    except Exception as e:
        print(f"\nERROR: Conversion failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
