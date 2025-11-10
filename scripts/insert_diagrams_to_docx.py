#!/usr/bin/env python3

"""
Insert rendered diagram images into Word document

This script reads the markdown file, finds Mermaid diagram blocks,
and replaces them with the corresponding PNG images in the Word document.
"""

import re
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx is required")
    print("Install with: pip install python-docx")
    print("\nNote: In NixOS, you may need to use a virtual environment:")
    print("  python3 -m venv venv")
    print("  source venv/bin/activate")
    print("  pip install python-docx")
    sys.exit(1)

def extract_sections_with_diagrams(markdown_file):
    """
    Parse markdown and identify sections with diagrams
    Returns list of sections with their content and diagram numbers
    """
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = []
    current_section = {'title': '', 'content': [], 'diagram_num': None}
    diagram_counter = 0
    in_mermaid = False

    for line in content.split('\n'):
        # Check for headers
        if line.startswith('#'):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'title': line.strip('#').strip(), 'content': [], 'diagram_num': None}

        # Check for mermaid block start
        if line.strip() == '```mermaid':
            in_mermaid = True
            diagram_counter += 1
            current_section['diagram_num'] = diagram_counter
            current_section['content'].append(f'<<DIAGRAM_{diagram_counter}>>')
            continue

        # Check for mermaid block end
        if in_mermaid and line.strip() == '```':
            in_mermaid = False
            continue

        # Skip mermaid diagram content
        if in_mermaid:
            continue

        # Add regular content
        if not in_mermaid:
            current_section['content'].append(line)

    # Add last section
    if current_section['content']:
        sections.append(current_section)

    return sections

def create_docx_with_diagrams(markdown_file, diagrams_dir, output_file):
    """Create Word document with diagrams inserted"""

    print("Creating Word document with diagrams...")

    # Create document
    doc = Document()

    # Set up styles
    styles = doc.styles
    style = styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Extract sections
    sections = extract_sections_with_diagrams(markdown_file)

    print(f"Processing {len(sections)} sections...")

    for section_idx, section in enumerate(sections, 1):
        # Add section title as heading
        if section['title']:
            # Determine heading level based on original markdown
            if section['title'] in ['GitHub-ServiceNow Data Integration in ARC']:
                doc.add_heading(section['title'], 0)
            else:
                # Count # characters to determine level
                level = 1
                for line in section['content']:
                    if line.startswith('#'):
                        level = len(line) - len(line.lstrip('#'))
                        break
                doc.add_heading(section['title'], min(level, 3))

        # Process content
        for line in section['content']:
            # Check if this is a diagram placeholder
            diagram_match = re.match(r'<<DIAGRAM_(\d+)>>', line.strip())
            if diagram_match:
                diagram_num = int(diagram_match.group(1))
                diagram_file = diagrams_dir / f'diagram-{diagram_num}.png'

                if diagram_file.exists():
                    # Add diagram caption
                    caption_para = doc.add_paragraph()
                    caption_run = caption_para.add_run(f'Diagram {diagram_num}')
                    caption_run.bold = True
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add diagram image
                    try:
                        picture = doc.add_picture(str(diagram_file), width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"  ✓ Inserted diagram {diagram_num}")
                    except Exception as e:
                        print(f"  ✗ Failed to insert diagram {diagram_num}: {e}")
                        doc.add_paragraph(f'[Diagram {diagram_num} - Image file error]')
                else:
                    # Diagram file doesn't exist
                    para = doc.add_paragraph(f'[Diagram {diagram_num} - Not rendered]')
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  ⚠ Diagram {diagram_num} file not found")

                # Add spacing
                doc.add_paragraph()
                continue

            # Regular content
            if line.strip():
                # Check for different markdown elements
                if line.startswith('```'):
                    # Code block
                    continue
                elif line.startswith('|') and '|' in line[1:]:
                    # Table row - skip for now (complex)
                    continue
                elif line.startswith('-') or line.startswith('*'):
                    # Bullet list
                    doc.add_paragraph(line.lstrip('-*').strip(), style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    # Numbered list
                    doc.add_paragraph(line.split('.', 1)[1].strip(), style='List Number')
                else:
                    # Regular paragraph
                    if line.strip():
                        doc.add_paragraph(line)

    # Save document
    doc.save(output_file)
    print(f"\n✓ Document saved: {output_file}")

    # Get file size
    file_size = os.path.getsize(output_file) / 1024  # KB
    print(f"  File size: {file_size:.1f} KB")

def update_existing_docx(input_docx, markdown_file, diagrams_dir, output_file):
    """
    Update existing Word document by finding mermaid code blocks and replacing with images
    """
    print("Updating existing Word document with diagrams...")

    # Load existing document
    doc = Document(input_docx)

    diagram_counter = 0
    paragraphs_to_process = []

    # Find paragraphs with mermaid code
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if 'graph TB' in text or 'sequenceDiagram' in text or 'erDiagram' in text:
            diagram_counter += 1
            paragraphs_to_process.append((i, diagram_counter))

    print(f"Found {diagram_counter} diagram code blocks in document")

    # Replace code blocks with images (process in reverse to maintain indices)
    for para_idx, diagram_num in reversed(paragraphs_to_process):
        diagram_file = diagrams_dir / f'diagram-{diagram_num}.png'

        # Clear the paragraph
        para = doc.paragraphs[para_idx]
        para.clear()

        if diagram_file.exists():
            # Add caption
            caption_run = para.add_run(f'Diagram {diagram_num}\n')
            caption_run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Insert image after the paragraph
            # Note: We need to insert in the correct position
            parent = para._element.getparent()
            para_element = para._element

            try:
                # Add picture in new paragraph
                new_para = para.insert_paragraph_before()
                run = new_para.add_run()
                run.add_picture(str(diagram_file), width=Inches(6.0))
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"  ✓ Inserted diagram {diagram_num}")
            except Exception as e:
                print(f"  ✗ Failed to insert diagram {diagram_num}: {e}")
                para.text = f'[Diagram {diagram_num} - Image file error]'
        else:
            para.text = f'[Diagram {diagram_num} - Not rendered]'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"  ⚠ Diagram {diagram_num} file not found")

    # Save updated document
    doc.save(output_file)
    print(f"\n✓ Document saved: {output_file}")

    # Get file size
    file_size = os.path.getsize(output_file) / 1024  # KB
    print(f"  File size: {file_size:.1f} KB")

def main():
    # Paths
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    docs_dir = project_root / 'docs'
    diagrams_dir = docs_dir / 'diagrams'

    markdown_file = docs_dir / 'GITHUB-SERVICENOW-DATA-Integration-in-ARC.md'
    input_docx = docs_dir / 'GITHUB-SERVICENOW-DATA-Integration-in-ARC.docx'
    output_file = docs_dir / 'GITHUB-SERVICENOW-DATA-Integration-in-ARC-with-diagrams.docx'

    print("=" * 70)
    print("Insert Diagrams into Word Document")
    print("=" * 70)
    print()

    # Check if files exist
    if not markdown_file.exists():
        print(f"ERROR: Markdown file not found: {markdown_file}")
        sys.exit(1)

    if not diagrams_dir.exists():
        print(f"ERROR: Diagrams directory not found: {diagrams_dir}")
        print("Run render_diagrams.py first to generate diagram images")
        sys.exit(1)

    # Check how many diagrams are available
    diagram_files = list(diagrams_dir.glob('diagram-*.png'))
    print(f"Found {len(diagram_files)} diagram images in {diagrams_dir}")
    print()

    if not diagram_files:
        print("ERROR: No diagram files found!")
        print("Run render_diagrams.py first to generate diagram images")
        sys.exit(1)

    # Update existing document
    if input_docx.exists():
        print(f"Input document: {input_docx}")
        print(f"Output document: {output_file}")
        print()
        update_existing_docx(input_docx, markdown_file, diagrams_dir, output_file)
    else:
        print(f"Input document not found, creating new document...")
        print(f"Output document: {output_file}")
        print()
        create_docx_with_diagrams(markdown_file, diagrams_dir, output_file)

    print()
    print("=" * 70)
    print("✓ Conversion completed!")
    print("=" * 70)
    print()
    print(f"Open the document with:")
    print(f"  libreoffice {output_file}")

if __name__ == '__main__':
    main()
